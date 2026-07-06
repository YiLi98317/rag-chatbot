#!/usr/bin/env python3
"""将 data/knowledge/ 写作模板 + 正式案例合并导出为 docx/txt。"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
KNOWLEDGE = ROOT / "data" / "knowledge"
TEMPLATES = KNOWLEDGE / "_templates"

# 按文档类型组织：模板 + 对应目录下的正式案例
TYPE_SECTIONS = [
    {
        "source_type": "sop",
        "label": "SOP 操作流程（办单 / 锁机 / 审核等）",
        "template": "sop_template.md",
        "case_dirs": ["02_sop"],
    },
    {
        "source_type": "faq",
        "label": "标准问答 FAQ",
        "template": "faq_template.md",
        "case_dirs": ["05_faq"],
    },
    {
        "source_type": "rule",
        "label": "平台规则 / 定价与费用",
        "template": "rule_template.md",
        "case_dirs": ["01_rules", "10_pricing_and_fee"],
    },
    {
        "source_type": "system_operation",
        "label": "系统页面操作",
        "template": "system_operation_template.md",
        "case_dirs": ["08_system_operations"],
    },
    {
        "source_type": "sales_script",
        "label": "销售话术",
        "template": "sales_script_template.md",
        "case_dirs": ["03_sales_scripts"],
    },
    {
        "source_type": "redline",
        "label": "红线与禁用话术",
        "template": "redline_template.md",
        "case_dirs": ["04_redlines"],
    },
    {
        "source_type": "after_sales",
        "label": "售后 / 质保 / 丢失 / 逾期",
        "template": "after_sales_template.md",
        "case_dirs": ["09_after_sales"],
    },
]

FRONTMATTER_RE = re.compile(r"^---\s*\n(.*?)\n---\s*\n", re.DOTALL)
HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
ORDERED_RE = re.compile(r"^(\d+)\.\s*(.*)$")
BULLET_RE = re.compile(r"^[-*]\s+(.*)$")
HTML_COMMENT_RE = re.compile(r"<!--.*?-->", re.DOTALL)


def parse_frontmatter(text: str) -> tuple[dict[str, str], str]:
    m = FRONTMATTER_RE.match(text)
    if not m:
        return {}, text
    meta: dict[str, str] = {}
    for line in m.group(1).splitlines():
        if ":" not in line:
            continue
        k, v = line.split(":", 1)
        meta[k.strip()] = v.strip()
    return meta, text[m.end():]


def load_md(path: Path) -> tuple[dict[str, str], str]:
    text = path.read_text(encoding="utf-8")
    meta, body = parse_frontmatter(text)
    body = HTML_COMMENT_RE.sub("", body).strip()
    return meta, body


def doc_sort_key(path: Path) -> tuple:
    m = re.search(r"(\d+)", path.stem)
    num = int(m.group(1)) if m else 9999
    return (num, path.name)


def collect_cases() -> dict[str, list[tuple[Path, dict[str, str], str]]]:
    by_type: dict[str, list[tuple[Path, dict[str, str], str]]] = {
        s["source_type"]: [] for s in TYPE_SECTIONS
    }
    seen: set[Path] = set()

    for section in TYPE_SECTIONS:
        st = section["source_type"]
        for dirname in section["case_dirs"]:
            d = KNOWLEDGE / dirname
            if not d.is_dir():
                continue
            for path in sorted(d.glob("*.md"), key=doc_sort_key):
                if path in seen:
                    continue
                meta, body = load_md(path)
                case_type = meta.get("source_type", st)
                bucket = case_type if case_type in by_type else st
                by_type[bucket].append((path, meta, body))
                seen.add(path)

    for st in by_type:
        by_type[st].sort(key=lambda x: doc_sort_key(x[0]))
    return by_type


def set_run_font(run, bold=False, italic=False, size=11, color=None):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "PingFang SC"
    if color:
        run.font.color.rgb = color


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=min(level, 3))
    for run in p.runs:
        run.font.name = "PingFang SC"


def add_body_paragraph(doc: Document, text: str, italic=False, quote=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, italic=italic or quote)
    if quote:
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    return p


def render_markdown_body(doc: Document, body: str):
    for line in body.splitlines():
        line = line.rstrip()
        if not line.strip():
            continue

        hm = HEADING_RE.match(line)
        if hm:
            add_heading(doc, hm.group(2).strip(), len(hm.group(1)))
            continue

        if line.startswith(">"):
            add_body_paragraph(doc, line.lstrip("> ").strip(), quote=True)
            continue

        om = ORDERED_RE.match(line)
        if om:
            p = doc.add_paragraph(style="List Number")
            set_run_font(p.add_run(om.group(2).strip()))
            continue

        bm = BULLET_RE.match(line)
        if bm:
            p = doc.add_paragraph(style="List Bullet")
            set_run_font(p.add_run(bm.group(1).strip()))
            continue

        if line.startswith(("Q：", "Q:")):
            p = doc.add_paragraph()
            set_run_font(p.add_run(line), bold=True)
            continue

        add_body_paragraph(doc, line)


def render_doc_block(doc: Document, path: Path, meta: dict, body: str, block_label: str):
    title = meta.get("title") or path.stem
    doc_id = meta.get("doc_id") or path.stem
    add_heading(doc, f"{block_label}：{doc_id} · {title}", 3)

    meta_bits = []
    for key in ("category", "visible_to", "risk_level", "owner", "version", "effective_from"):
        val = meta.get(key, "")
        if val and not val.startswith("#"):
            meta_bits.append(f"{key}：{val}")
    if meta_bits:
        p = doc.add_paragraph()
        set_run_font(p.add_run("｜".join(meta_bits)), italic=True, size=9, color=RGBColor(0x88, 0x88, 0x88))

    render_markdown_body(doc, body)
    doc.add_paragraph()


def build_document(cases_by_type: dict) -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "PingFang SC"
    style.font.size = Pt(11)

    template_count = len(TYPE_SECTIONS)
    case_count = sum(len(v) for v in cases_by_type.values())

    add_heading(doc, "手机妈妈 RAG 知识库（模板 + 正式案例）", 0)
    doc.add_paragraph(
        "说明：本文件合并了 data/knowledge/_templates/ 下的写作模板，"
        "以及各目录中的正式案例文档。\n"
        "每一类文档的结构是：先看【写作模板】了解该怎么写，再看【正式案例】了解已入库的内容。\n"
        "这不是简单的「问题 + 答案」对照表，SOP 类包含完整操作流程，FAQ 类才有标准问答结构。"
    )
    p = doc.add_paragraph()
    set_run_font(p.add_run(f"模板 {template_count} 份｜正式案例 {case_count} 篇"), bold=True)

    add_heading(doc, "目录概览", 1)
    for i, section in enumerate(TYPE_SECTIONS, 1):
        st = section["source_type"]
        n = len(cases_by_type.get(st, []))
        p = doc.add_paragraph(style="List Number")
        suffix = f"（模板 1 份 + 案例 {n} 篇）" if n else "（仅模板，暂无正式案例）"
        set_run_font(p.add_run(f"{section['label']}{suffix}"))

    doc.add_page_break()

    for i, section in enumerate(TYPE_SECTIONS, 1):
        st = section["source_type"]
        label = section["label"]
        add_heading(doc, f"{i}. {label}", 1)

        tpl_path = TEMPLATES / section["template"]
        if tpl_path.is_file():
            meta, body = load_md(tpl_path)
            render_doc_block(doc, tpl_path, meta, body, "【写作模板】")
        else:
            add_body_paragraph(doc, f"（未找到模板文件：{section['template']}）", italic=True)

        cases = cases_by_type.get(st, [])
        if cases:
            add_heading(doc, f"【正式案例】共 {len(cases)} 篇", 2)
            for path, meta, body in cases:
                render_doc_block(doc, path, meta, body, "案例")
        else:
            add_heading(doc, "【正式案例】暂无", 2)
            add_body_paragraph(doc, "该类型目前只有写作模板，尚未入库正式案例。", italic=True)

        if i < len(TYPE_SECTIONS):
            doc.add_page_break()

    return doc


def build_txt(cases_by_type: dict) -> str:
    lines = [
        "手机妈妈 RAG 知识库（模板 + 正式案例）",
        "=" * 50,
        "",
        "说明：每类文档 = 写作模板 + 正式案例。",
        "",
    ]
    for i, section in enumerate(TYPE_SECTIONS, 1):
        st = section["source_type"]
        lines += ["", "=" * 50, f"{i}. {section['label']}", "=" * 50, ""]

        tpl_path = TEMPLATES / section["template"]
        if tpl_path.is_file():
            _, body = load_md(tpl_path)
            lines += [f"--- 【写作模板】{section['template']} ---", body, ""]

        cases = cases_by_type.get(st, [])
        if cases:
            lines.append(f"--- 【正式案例】共 {len(cases)} 篇 ---")
            for path, meta, body in cases:
                title = meta.get("title") or path.stem
                doc_id = meta.get("doc_id") or path.stem
                lines += ["", f"### 案例：{doc_id} · {title}", body, ""]
        else:
            lines.append("（暂无正式案例）")

    return "\n".join(lines)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        default=ROOT / "business_review" / "手机妈妈知识库_模板与案例.docx",
    )
    parser.add_argument("--txt", action="store_true", help="同时输出 txt")
    args = parser.parse_args()

    cases_by_type = collect_cases()
    case_total = sum(len(v) for v in cases_by_type.values())

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document(cases_by_type)
    doc.save(str(args.output))
    print(f"已导出 docx: {args.output}")
    print(f"  模板 {len(TYPE_SECTIONS)} 份 + 正式案例 {case_total} 篇")

    if args.txt:
        txt_path = args.output.with_suffix(".txt")
        txt_path.write_text(build_txt(cases_by_type), encoding="utf-8")
        print(f"已导出 txt:  {txt_path}")


if __name__ == "__main__":
    main()
